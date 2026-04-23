#!/usr/bin/env python3
"""
Word Document to Markdown Converter
Converts .docx files to Markdown format for book publishing.

Usage:
    python3 docx_to_markdown.py input.docx [-o output.md] [--extract-images]

Features:
- Preserves headings (H1-H6) as Markdown headings
- Converts bold, italic, underline formatting
- Handles bullet and numbered lists
- Converts tables to Markdown format
- Optionally extracts embedded images

Dependencies: pip3 install python-docx
"""

import argparse
import re
import os
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed.")
    print("Install with: pip3 install python-docx")
    exit(1)


def extract_paragraph_text(paragraph, preserve_formatting=True):
    """Extract text from a paragraph with optional formatting preservation."""
    if not preserve_formatting:
        return paragraph.text

    result = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        result.append(text)

    return ''.join(result)


def heading_level_from_style(style_name):
    """Determine heading level from Word style name."""
    if not style_name:
        return 0

    style_lower = style_name.lower()

    # Check for explicit heading styles
    if 'heading 1' in style_lower or style_lower == 'heading1':
        return 1
    if 'heading 2' in style_lower or style_lower == 'heading2':
        return 2
    if 'heading 3' in style_lower or style_lower == 'heading3':
        return 3
    if 'heading 4' in style_lower or style_lower == 'heading4':
        return 4
    if 'heading 5' in style_lower or style_lower == 'heading5':
        return 5
    if 'heading 6' in style_lower or style_lower == 'heading6':
        return 6
    if 'title' in style_lower:
        return 1
    if 'subtitle' in style_lower:
        return 2

    return 0


def is_list_paragraph(paragraph):
    """Check if paragraph is a list item."""
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            return True
    return False


def get_list_level(paragraph):
    """Get the indentation level of a list item."""
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            ilvl = numPr.ilvl
            if ilvl is not None:
                return ilvl.val
    return 0


def convert_table(table):
    """Convert a Word table to Markdown format."""
    lines = []

    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get cell text, handle merged cells
            cell_text = cell.text.strip().replace('\n', ' ')
            cells.append(cell_text)

        lines.append('| ' + ' | '.join(cells) + ' |')

        # Add separator after header row
        if i == 0:
            separator = '| ' + ' | '.join(['---' for _ in cells]) + ' |'
            lines.append(separator)

    return '\n'.join(lines)


def extract_images(doc, output_dir):
    """Extract images from document and save to output directory."""
    images_dir = output_dir / 'images'
    images_dir.mkdir(exist_ok=True)

    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_ext = rel.target_ref.split('.')[-1]
            image_name = f"image_{image_count:03d}.{image_ext}"
            image_path = images_dir / image_name

            with open(image_path, 'wb') as f:
                f.write(rel.target_part.blob)

            print(f"  Extracted: {image_name}")

    return image_count


def convert_docx_to_markdown(input_path, output_path=None, extract_imgs=False):
    """Convert a Word document to Markdown."""
    input_path = Path(input_path)

    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return None

    if output_path:
        output_path = Path(output_path)
    else:
        output_path = input_path.with_suffix('.md')

    print(f"Converting: {input_path.name}")
    print(f"Output: {output_path}")

    # Load document
    doc = Document(str(input_path))

    # Extract images if requested
    if extract_imgs:
        img_count = extract_images(doc, output_path.parent)
        print(f"  Extracted {img_count} images")

    # Convert content
    markdown_lines = []
    in_list = False
    list_type = None  # 'bullet' or 'number'

    for element in doc.element.body:
        # Handle tables
        if element.tag.endswith('tbl'):
            # Find the table object
            for table in doc.tables:
                if table._tbl == element:
                    if in_list:
                        markdown_lines.append('')
                        in_list = False
                    markdown_lines.append('')
                    markdown_lines.append(convert_table(table))
                    markdown_lines.append('')
                    break
            continue

        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the paragraph object
            for para in doc.paragraphs:
                if para._p == element:
                    text = extract_paragraph_text(para)
                    style_name = para.style.name if para.style else ''
                    heading_level = heading_level_from_style(style_name)

                    # Skip empty paragraphs (but preserve paragraph breaks)
                    if not text.strip():
                        if in_list:
                            markdown_lines.append('')
                            in_list = False
                        markdown_lines.append('')
                        break

                    # Handle headings
                    if heading_level > 0:
                        if in_list:
                            markdown_lines.append('')
                            in_list = False
                        prefix = '#' * heading_level
                        markdown_lines.append('')
                        markdown_lines.append(f"{prefix} {text}")
                        markdown_lines.append('')
                        break

                    # Handle lists
                    if is_list_paragraph(para):
                        level = get_list_level(para)
                        indent = '  ' * level

                        # Detect list type from style
                        if 'bullet' in style_name.lower() or 'list bullet' in style_name.lower():
                            markdown_lines.append(f"{indent}- {text}")
                        elif 'number' in style_name.lower() or 'list number' in style_name.lower():
                            markdown_lines.append(f"{indent}1. {text}")
                        else:
                            # Default to bullet
                            markdown_lines.append(f"{indent}- {text}")

                        in_list = True
                        break

                    # Handle blockquotes (based on style or indentation)
                    if 'quote' in style_name.lower() or 'block' in style_name.lower():
                        if in_list:
                            markdown_lines.append('')
                            in_list = False
                        markdown_lines.append(f"> {text}")
                        break

                    # Regular paragraph
                    if in_list:
                        markdown_lines.append('')
                        in_list = False
                    markdown_lines.append(text)
                    break

    # Clean up excessive blank lines
    content = '\n'.join(markdown_lines)
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = content.strip() + '\n'

    # Write output
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

    # Count elements
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    table_count = len(doc.tables)

    print(f"  Paragraphs: {para_count}")
    print(f"  Tables: {table_count}")
    print(f"  Output size: {len(content)} characters")

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word documents (.docx) to Markdown format'
    )
    parser.add_argument(
        'input',
        help='Input .docx file path'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output .md file path (default: same name as input)'
    )
    parser.add_argument(
        '--extract-images',
        action='store_true',
        help='Extract embedded images to images/ subdirectory'
    )

    args = parser.parse_args()

    result = convert_docx_to_markdown(
        args.input,
        args.output,
        args.extract_images
    )

    if result:
        print(f"\nConversion complete: {result}")
    else:
        print("\nConversion failed")
        exit(1)


if __name__ == '__main__':
    main()
